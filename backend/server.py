#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
╔════════════════════════════════════════════════════════════╗
║  🎯 TRANSCRITOR 2027 — Backend Flask                       ║
║  API REST para transcrição com Whisper                     ║
╚════════════════════════════════════════════════════════════╝
"""

import os
import uuid
import subprocess
import json
from datetime import datetime
from pathlib import Path
from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
from faster_whisper import WhisperModel
from docx import Document
from fpdf import FPDF

# ═══════════════════════════════════════════════════════════════
# CONFIGURAÇÕES
# ═══════════════════════════════════════════════════════════════

app = Flask(__name__)
CORS(app)  # Habilita CORS para frontend

# Pastas
BASE_DIR = Path(__file__).parent.parent
UPLOAD_DIR = BASE_DIR / 'uploads'
DOWNLOAD_DIR = BASE_DIR / 'downloads'

UPLOAD_DIR.mkdir(exist_ok=True)
DOWNLOAD_DIR.mkdir(exist_ok=True)

# Limites
MAX_FILE_SIZE = 2 * 1024 * 1024 * 1024  # 2GB
ALLOWED_EXTENSIONS = {'mp4', 'mov', 'avi', 'mkv', 'mp3', 'wav', 'm4a', 'webm', 'flac'}

# Modelo Whisper (carrega uma vez)
print("🧠 Carregando modelo Whisper...")
whisper_model = WhisperModel("small", device="cpu", compute_type="int8")
print("✅ Modelo carregado!")

# Armazenamento de estado (em produção use Redis/Database)
transcription_jobs = {}

# ═══════════════════════════════════════════════════════════════
# UTILITÁRIOS
# ═══════════════════════════════════════════════════════════════

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_audio(video_path, audio_path):
    """Extrai áudio usando FFmpeg"""
    cmd = [
        'ffmpeg',
        '-i', str(video_path),
        '-vn',
        '-acodec', 'pcm_s16le',
        '-ar', '16000',
        '-ac', '1',
        '-y',
        str(audio_path)
    ]
    subprocess.run(cmd, check=True, capture_output=True)

def transcribe_audio(audio_path, model_size='small', language='pt'):
    """Transcreve áudio com Whisper"""
    model = WhisperModel(model_size, device="cpu", compute_type="int8")
    
    segments, info = model.transcribe(
        str(audio_path),
        language=language if language != 'auto' else None,
        beam_size=5,
        best_of=5,
        temperature=0.2,
        vad_filter=True
    )
    
    text = " ".join(seg.text.strip() for seg in segments).strip()
    
    return {
        'text': text,
        'language': info.language,
        'language_probability': info.language_probability,
        'duration': info.duration
    }

def save_txt(text, output_path):
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(f"Transcrição gerada em: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n")
        f.write("=" * 60 + "\n\n")
        f.write(text)

def save_docx(text, metadata, output_path):
    doc = Document()
    doc.add_heading('📋 Transcrição da Reunião', level=1)
    doc.add_paragraph(f"Arquivo: {metadata.get('filename', 'N/A')}")
    doc.add_paragraph(f"Gerado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph(f"Modelo: {metadata.get('model', 'small')}")
    doc.add_paragraph(f"Idioma: {metadata.get('language', 'pt').upper()}")
    doc.add_paragraph("-" * 50)
    doc.add_paragraph(text)
    doc.save(output_path)

def save_pdf(text, output_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    pdf.cell(200, 10, txt="Transcrição da Reunião", ln=True, align="C")
    pdf.ln(5)
    texto_pdf = text.encode('latin-1', errors='replace').decode('latin-1')
    pdf.multi_cell(0, 5, txt=texto_pdf)
    pdf.output(output_path)

# ═══════════════════════════════════════════════════════════════
# ROTAS DA API
# ═══════════════════════════════════════════════════════════════

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Recebe arquivo de áudio/vídeo"""
    if 'file' not in request.files:
        return jsonify({'error': 'Nenhum arquivo enviado'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'Nome de arquivo vazio'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Tipo de arquivo não permitido'}), 400
    
    # Salvar arquivo
    file_id = str(uuid.uuid4())
    filename = secure_filename(file.filename)
    file_path = UPLOAD_DIR / f"{file_id}_{filename}"
    
    file.save(str(file_path))
    
    # Armazenar metadados
    transcription_jobs[file_id] = {
        'filename': filename,
        'file_path': str(file_path),
        'status': 'uploaded',
        'created_at': datetime.now().isoformat()
    }
    
    return jsonify({
        'file_id': file_id,
        'filename': filename,
        'size': file_path.stat().st_size
    })

@app.route('/api/transcribe', methods=['POST'])
def transcribe():
    """Inicia transcrição"""
    data = request.json
    file_id = data.get('file_id')
    model = data.get('model', 'small')
    language = data.get('language', 'pt')
    output_format = data.get('format', 'txt')
    
    if not file_id or file_id not in transcription_jobs:
        return jsonify({'error': 'Arquivo não encontrado'}), 404
    
    job = transcription_jobs[file_id]
    
    try:
        # Extrair áudio
        audio_path = DOWNLOAD_DIR / f"{file_id}.wav"
        extract_audio(job['file_path'], audio_path)
        
        # Transcrever
        result = transcribe_audio(audio_path, model, language)
        
        # Salvar em diferentes formatos
        download_urls = {}
        
        # TXT
        txt_path = DOWNLOAD_DIR / f"{file_id}.txt"
        save_txt(result['text'], txt_path)
        download_urls['txt'] = f"/api/download/{file_id}/txt"
        
        # DOCX
        if output_format in ['docx', 'all']:
            docx_path = DOWNLOAD_DIR / f"{file_id}.docx"
            save_docx(result['text'], {
                'filename': job['filename'],
                'model': model,
                'language': result['language']
            }, docx_path)
            download_urls['docx'] = f"/api/download/{file_id}/docx"
        
        # PDF
        if output_format in ['pdf', 'all']:
            pdf_path = DOWNLOAD_DIR / f"{file_id}.pdf"
            save_pdf(result['text'], pdf_path)
            download_urls['pdf'] = f"/api/download/{file_id}/pdf"
        
        # Atualizar job
        job['status'] = 'completed'
        job['result'] = result
        job['download_urls'] = download_urls
        job['audio_path'] = str(audio_path)
        
        return jsonify({
            'file_id': file_id,
            'text': result['text'],
            'language': result['language'],
            'duration': result['duration'],
            'download_urls': download_urls,
            'audio_url': f"/api/audio/{file_id}"
        })
        
    except Exception as e:
        job['status'] = 'failed'
        job['error'] = str(e)
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/<file_id>/<format>', methods=['GET'])
def download(file_id, format):
    """Baixa arquivo transcrito"""
    if file_id not in transcription_jobs:
        return jsonify({'error': 'Job não encontrado'}), 404
    
    file_path = DOWNLOAD_DIR / f"{file_id}.{format}"
    
    if not file_path.exists():
        return jsonify({'error': 'Arquivo não encontrado'}), 404
    
    return send_file(
        file_path,
        as_attachment=True,
        download_name=f"transcricao_{file_id}.{format}"
    )

@app.route('/api/audio/<file_id>', methods=['GET'])
def get_audio(file_id):
    """Baixa áudio extraído"""
    if file_id not in transcription_jobs:
        return jsonify({'error': 'Job não encontrado'}), 404
    
    job = transcription_jobs[file_id]
    audio_path = job.get('audio_path')
    
    if not audio_path or not Path(audio_path).exists():
        return jsonify({'error': 'Áudio não encontrado'}), 404
    
    return send_file(audio_path, as_attachment=True, download_name=f"audio_{file_id}.wav")

@app.route('/api/status/<file_id>', methods=['GET'])
def get_status(file_id):
    """Verifica status do job"""
    if file_id not in transcription_jobs:
        return jsonify({'error': 'Job não encontrado'}), 404
    
    job = transcription_jobs[file_id]
    return jsonify({
        'file_id': file_id,
        'status': job['status'],
        'error': job.get('error')
    })

@app.route('/api/jobs', methods=['GET'])
def list_jobs():
    """Lista todos os jobs (para debug)"""
    return jsonify({
        file_id: {
            'filename': job['filename'],
            'status': job['status'],
            'created_at': job['created_at']
        }
        for file_id, job in transcription_jobs.items()
    })

# ═══════════════════════════════════════════════════════════════
# SERVIDOR
# ═══════════════════════════════════════════════════════════════

if __name__ == '__main__':
    print("🚀 Iniciando servidor Transcritor 2027...")
    print(f"📁 Uploads: {UPLOAD_DIR}")
    print(f"📁 Downloads: {DOWNLOAD_DIR}")
    print("🌐 API: http://localhost:5000/api")
    app.run(debug=True, host='0.0.0.0', port=5000)
